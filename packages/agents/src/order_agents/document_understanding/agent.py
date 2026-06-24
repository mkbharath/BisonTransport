"""Document Understanding Agent — extracts text and tables from attachments.

Processes PDFs (digital + scanned), images, Excel, and Word documents.
Publishes combined extraction results to the extraction queue.
"""

import json
import uuid
from typing import Any

from sqlalchemy import text

from order_shared.adapters import get_adapters
from order_shared.adapters.base import QueueMessage
from order_shared.db.session import async_session_factory
from order_shared.models.enums import AgentType
from order_shared.utils.logger import get_logger
from order_agents.base_agent import BaseAgent

logger = get_logger(__name__)


class DocumentUnderstandingAgent(BaseAgent):
    """Processes email attachments to extract text and structured data."""

    agent_type = AgentType.DOCUMENT_UNDERSTANDING
    input_queue = "document-processing"

    async def process_message(self, message: QueueMessage) -> None:
        """Process all attachments for an email."""
        adapters = get_adapters()
        email_id = message.body["email_id"]
        attachment_ids = message.body.get("attachment_ids", [])

        # Get email body text from DB
        async with async_session_factory() as session:
            result = await session.execute(
                text("SELECT body_text, body_html FROM emails WHERE id = :id"),
                {"id": email_id},
            )
            row = result.fetchone()
            email_body = row[0] or row[1] or "" if row else ""

        # Process each attachment
        extraction_results: list[dict[str, Any]] = []

        for att_id in attachment_ids:
            async with async_session_factory() as session:
                result = await session.execute(
                    text("SELECT file_name, file_type, s3_key FROM email_attachments WHERE id = :id"),
                    {"id": att_id},
                )
                att_row = result.fetchone()
                if not att_row:
                    continue

                file_name, file_type, s3_key = att_row

            # Download file from storage
            file_data = await adapters.storage.download_file(bucket="attachments", key=s3_key)

            # Extract based on file type
            extracted = await self._extract_document(file_data, file_type, file_name)

            # Store extracted text to storage
            extracted_key = f"extracted-text/{email_id}/{att_id}.json"
            await adapters.storage.upload_file(
                bucket="extracted-text",
                key=extracted_key,
                data=json.dumps(extracted).encode(),
                content_type="application/json",
            )

            # Update attachment record
            async with async_session_factory() as session:
                await session.execute(
                    text("""
                        UPDATE email_attachments
                        SET extracted_text_s3_key = :key, ocr_confidence = :conf,
                            processing_status = 'completed'
                        WHERE id = :id
                    """),
                    {"key": extracted_key, "conf": extracted.get("confidence", 0), "id": att_id},
                )
                await session.commit()

            extraction_results.append(extracted)
            self.logger.info(
                f"Extracted {file_name}: {len(extracted.get('text', ''))} chars, "
                f"confidence={extracted.get('confidence', 0):.1f}%"
            )

        # Publish combined results to extraction queue
        await adapters.queue.publish_message(
            queue_name="extraction",
            body={
                "email_id": email_id,
                "email_body": email_body,
                "extraction_results": extraction_results,
                "attachment_count": len(extraction_results),
                "is_customer_response": message.body.get("is_customer_response", False),
                "related_order_id": message.body.get("related_order_id"),
            },
        )

    async def _extract_document(
        self, file_data: bytes, file_type: str, file_name: str
    ) -> dict[str, Any]:
        """Extract text and tables from a document based on its type."""
        adapters = get_adapters()
        file_type_lower = (file_type or "").lower().strip(".")

        if file_type_lower == "pdf":
            return await self._extract_pdf(file_data, file_name)
        elif file_type_lower in ("png", "jpg", "jpeg", "tiff", "bmp"):
            return await self._extract_image(file_data, file_type_lower, file_name)
        elif file_type_lower in ("xlsx", "xls"):
            return await self._extract_excel(file_data, file_name)
        elif file_type_lower in ("docx", "doc"):
            return await self._extract_word(file_data, file_name)
        else:
            self.logger.warning(f"Unsupported file type: {file_type_lower} for {file_name}")
            return {"text": "", "confidence": 0, "file_name": file_name, "tables": []}

    async def _extract_pdf(self, file_data: bytes, file_name: str) -> dict[str, Any]:
        """Extract from PDF using pdfplumber (digital) or OCR (scanned)."""
        adapters = get_adapters()
        ocr_result = await adapters.ocr.extract_text(file_data, "pdf")
        tables = await adapters.ocr.extract_tables(file_data, "pdf")

        return {
            "text": ocr_result.text,
            "confidence": ocr_result.confidence,
            "pages": ocr_result.pages,
            "tables": tables,
            "file_name": file_name,
            "file_type": "pdf",
        }

    async def _extract_image(
        self, file_data: bytes, file_type: str, file_name: str
    ) -> dict[str, Any]:
        """Extract text from image using OCR."""
        adapters = get_adapters()
        ocr_result = await adapters.ocr.extract_text(file_data, file_type)

        return {
            "text": ocr_result.text,
            "confidence": ocr_result.confidence,
            "pages": ocr_result.pages,
            "tables": [],
            "file_name": file_name,
            "file_type": file_type,
        }

    async def _extract_excel(self, file_data: bytes, file_name: str) -> dict[str, Any]:
        """Extract data from Excel files using openpyxl."""
        import io
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
        all_text_parts: list[str] = []
        tables: list[dict[str, Any]] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            # Use first row as headers
            headers = [str(cell) if cell else f"col_{i}" for i, cell in enumerate(rows[0])]
            table_rows = []
            for row in rows[1:]:
                row_dict = {}
                row_text_parts = []
                for i, cell in enumerate(row):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    value = str(cell) if cell is not None else ""
                    row_dict[key] = value
                    if value:
                        row_text_parts.append(f"{key}: {value}")
                table_rows.append(row_dict)
                all_text_parts.append(", ".join(row_text_parts))

            tables.append({
                "sheet": sheet_name,
                "headers": headers,
                "rows": table_rows,
            })

        return {
            "text": "\n".join(all_text_parts),
            "confidence": 90.0,  # Excel data is structured, high confidence
            "tables": tables,
            "file_name": file_name,
            "file_type": "xlsx",
        }

    async def _extract_word(self, file_data: bytes, file_name: str) -> dict[str, Any]:
        """Extract text from Word documents using python-docx."""
        import io
        import docx

        doc = docx.Document(io.BytesIO(file_data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text_content = "\n".join(paragraphs)

        # Extract tables
        tables: list[dict[str, Any]] = []
        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                row_dict = {
                    headers[i] if i < len(headers) else f"col_{i}": cell.text.strip()
                    for i, cell in enumerate(row.cells)
                }
                rows.append(row_dict)
            tables.append({"table_index": table_idx, "headers": headers, "rows": rows})

        return {
            "text": text_content,
            "confidence": 92.0,  # Word text is reliable
            "tables": tables,
            "file_name": file_name,
            "file_type": "docx",
        }
