"""Generate Excel and Image test attachments for all document processing scenarios."""

import os

OUTPUT_DIR = "test-emails/fixtures"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_excel():
    """Generate an Excel file with shipment order details."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()

    # Sheet 1: Order Details
    ws = wb.active
    ws.title = "Order Details"

    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    # Title
    ws.merge_cells("A1:B1")
    ws["A1"] = "TRANSPORTATION ORDER"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A2:B2")
    ws["A2"] = "Order Date: June 29, 2026"
    ws["A2"].font = Font(size=10, italic=True)

    row = 4
    sections = [
        ("CUSTOMER INFORMATION", [
            ("Customer Name", "Bharat Electronics Ltd"),
            ("Contact Person", "Vikram Singh"),
            ("Email", "bharathm@ideyalabs.com"),
            ("Phone", "+91 9123456789"),
            ("PO Number", "BEL-2026-9922"),
        ]),
        ("PICKUP DETAILS", [
            ("Location Name", "BEL Manufacturing Unit"),
            ("Address", "45 Jalahalli Post, Bangalore, KA 560013, India"),
            ("Date", "August 20, 2026"),
            ("Time Window", "06:00 - 09:00"),
            ("Instructions", "Security clearance required. Government facility."),
        ]),
        ("DELIVERY DETAILS", [
            ("Location Name", "BEL Hyderabad Office"),
            ("Address", "16 Nacharam Industrial Area, Hyderabad, TS 500076, India"),
            ("Date", "August 22, 2026"),
            ("Time Window", "10:00 - 14:00"),
            ("Instructions", "Deliver to receiving dock B. Call before arrival."),
        ]),
        ("SHIPMENT DETAILS", [
            ("Commodity", "Electronic defense equipment and radar components"),
            ("Freight Type", "FTL"),
            ("Total Weight", "32,000 lbs"),
            ("Number of Pallets", "18"),
            ("Stackable", "No"),
            ("Equipment Type", "Dry Van"),
            ("Truck Size", "53 ft"),
            ("Hazmat", "No"),
            ("Special Handling", "High-value cargo. GPS tracking required."),
        ]),
    ]

    for section_title, fields in sections:
        ws.cell(row=row, column=1, value=section_title).font = Font(bold=True, size=11, color="1F4E79")
        row += 1
        for label, value in fields:
            ws.cell(row=row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row, column=2, value=value)
            row += 1
        row += 1

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 55

    # Sheet 2: Table Format
    ws2 = wb.create_sheet("Shipment Table")
    headers = ["Field", "Value"]
    for col, h in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    table_data = [
        ("Customer", "Bharat Electronics Ltd"),
        ("Contact", "Vikram Singh"),
        ("Email", "bharathm@ideyalabs.com"),
        ("Phone", "+91 9123456789"),
        ("PO Number", "BEL-2026-9922"),
        ("Pickup Location", "BEL Manufacturing Unit, 45 Jalahalli Post, Bangalore, KA 560013"),
        ("Pickup Date", "2026-08-20"),
        ("Pickup Time", "06:00 - 09:00"),
        ("Delivery Location", "BEL Hyderabad Office, 16 Nacharam Industrial Area, Hyderabad, TS 500076"),
        ("Delivery Date", "2026-08-22"),
        ("Delivery Time", "10:00 - 14:00"),
        ("Commodity", "Electronic defense equipment and radar components"),
        ("Freight Type", "FTL"),
        ("Weight (lbs)", "32000"),
        ("Pallets", "18"),
        ("Equipment", "Dry Van"),
        ("Hazmat", "No"),
        ("Notes", "High-value cargo. GPS tracking required."),
    ]

    for i, (field, value) in enumerate(table_data, 2):
        ws2.cell(row=i, column=1, value=field).border = thin_border
        ws2.cell(row=i, column=2, value=value).border = thin_border

    ws2.column_dimensions["A"].width = 20
    ws2.column_dimensions["B"].width = 65

    filepath = os.path.join(OUTPUT_DIR, "scenario-excel-order.xlsx")
    wb.save(filepath)
    print(f"  Excel: {filepath}")


def generate_image():
    """Generate a PNG image simulating a printed/scanned order form."""
    from PIL import Image, ImageDraw, ImageFont

    # Create white image
    width, height = 800, 1000
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    # Use default font (no external font needed)
    try:
        font_large = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 24)
        font_medium = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
        font_small = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
    except (IOError, OSError):
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()

    y = 30

    # Header
    draw.text((250, y), "SHIPMENT ORDER", fill="black", font=font_large)
    y += 40
    draw.text((300, y), "Form #: SO-2026-4455", fill="gray", font=font_small)
    y += 30
    draw.line([(30, y), (770, y)], fill="black", width=2)
    y += 20

    # Order details as text
    lines = [
        ("CUSTOMER:", ""),
        ("  Customer Name:", "Mumbai Freight Corp"),
        ("  Contact:", "Anita Desai"),
        ("  Email:", "bharathm@ideyalabs.com"),
        ("  Phone:", "+91 9876012345"),
        ("", ""),
        ("PICKUP:", ""),
        ("  Location:", "Mumbai Freight Warehouse"),
        ("  Address:", "88 Andheri East, Mumbai, MH 400069, India"),
        ("  Date:", "September 1, 2026"),
        ("  Time:", "07:00 - 10:00"),
        ("  Notes:", "Gate pass required for entry"),
        ("", ""),
        ("DELIVERY:", ""),
        ("  Location:", "MFC Pune Hub"),
        ("  Address:", "22 Hinjewadi IT Park, Pune, MH 411057, India"),
        ("  Date:", "September 2, 2026"),
        ("  Time:", "11:00 - 15:00"),
        ("  Notes:", "Dock 3, call warehouse manager"),
        ("", ""),
        ("SHIPMENT:", ""),
        ("  PO Number:", "MFC-2026-3344"),
        ("  Commodity:", "Packaged food products and beverages"),
        ("  Freight Type:", "FTL"),
        ("  Weight:", "28,000 lbs"),
        ("  Pallets:", "16"),
        ("  Stackable:", "Yes"),
        ("  Equipment:", "Dry Van"),
        ("  Hazmat:", "No"),
        ("  Special:", "Temperature sensitive - keep below 25C"),
    ]

    for label, value in lines:
        if label == "":
            y += 10
            continue
        if label.startswith("  "):
            draw.text((50, y), label, fill="black", font=font_small)
            draw.text((250, y), value, fill="black", font=font_small)
        else:
            draw.text((30, y), label, fill="navy", font=font_medium)
            draw.text((250, y), value, fill="black", font=font_medium)
        y += 22

    # Footer line
    y += 20
    draw.line([(30, y), (770, y)], fill="black", width=1)
    y += 10
    draw.text((250, y), "Mumbai Freight Corp - Est. 1995", fill="gray", font=font_small)

    filepath = os.path.join(OUTPUT_DIR, "scenario-image-order.png")
    img.save(filepath)
    print(f"  Image: {filepath}")


def generate_word():
    """Generate a Word document with shipment order details."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("TRANSPORTATION ORDER", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Date: June 29, 2026 | Ref: WO-2026-7711").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Sections
    sections = [
        ("Customer Information", [
            ("Customer", "Reliance Retail Ltd"),
            ("Contact", "Sanjay Mehta"),
            ("Email", "bharathm@ideyalabs.com"),
            ("Phone", "+91 9988001122"),
        ]),
        ("Pickup Details", [
            ("Location", "Reliance DC Navi Mumbai"),
            ("Address", "Plot 5, TTC MIDC, Navi Mumbai, MH 400710, India"),
            ("Date", "August 25, 2026"),
            ("Time", "05:00 - 08:00"),
            ("Instructions", "Use rear gate. Night shift security will assist."),
        ]),
        ("Delivery Details", [
            ("Location", "Reliance Fresh Store #412"),
            ("Address", "MG Road, Camp Area, Pune, MH 411001, India"),
            ("Date", "August 26, 2026"),
            ("Time", "06:00 - 09:00"),
            ("Instructions", "Basement parking entrance for unloading"),
        ]),
        ("Shipment Details", [
            ("PO Number", "RRL-2026-55678"),
            ("Commodity", "FMCG products - packaged snacks and beverages"),
            ("Freight Type", "FTL"),
            ("Weight", "26,000 lbs"),
            ("Pallets", "22"),
            ("Stackable", "Yes"),
            ("Equipment", "Reefer"),
            ("Temperature", "2C to 8C"),
            ("Hazmat", "No"),
            ("Special Handling", "Maintain cold chain. Temperature logger required."),
        ]),
    ]

    for section_title, fields in sections:
        heading = doc.add_heading(section_title, level=2)
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"

        for i, (label, value) in enumerate(fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            # Bold the label
            for paragraph in table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph("")

    filepath = os.path.join(OUTPUT_DIR, "scenario-word-order.docx")
    doc.save(filepath)
    print(f"  Word: {filepath}")


if __name__ == "__main__":
    print("Generating test attachments...")
    print()

    print("1. Excel (XLSX):")
    generate_excel()

    print("2. Image (PNG):")
    generate_image()

    print("3. Word (DOCX):")
    try:
        generate_word()
    except ImportError:
        print("  Skipped (python-docx not installed locally). Run: pip3 install python-docx")

    print()
    print("Files generated in test-emails/fixtures/")
    print()
    print("To test, send an email to iltransport@ideyalabs.com with:")
    print("  Subject: 'Shipment Order - see attached'")
    print("  Body: 'Please process the attached order.'")
    print("  Attach: one of the generated files")
